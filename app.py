import streamlit as st
from pathlib import Path
from docx import Document  # to read Word (.docx) files
import re                  # clause regex
import pandas as pd        # displaying clause table
import math                # for cosine similarity
from datetime import datetime  # for history timestamps

# Try to import pdfplumber, but don't crash if it's not installed
try:
    import pdfplumber      # for reading PDF tenders
    PDF_SUPPORT = True
except ImportError:
    pdfplumber = None
    PDF_SUPPORT = False

# New OpenAI client (SDK v1+)
try:
    from openai import OpenAI
    OPENAI_AVAILABLE = True
except ImportError:
    OpenAI = None
    OPENAI_AVAILABLE = False


# ---------- BASIC PAGE SETTINGS (MUST BE FIRST STREAMLIT CALL) ----------
st.set_page_config(
    page_title="🐼 Wavetec Tender Library",
    layout="wide"
)

# ---------- GLOBAL STYLE OVERRIDES FOR CLEANER TEXT ----------
CUSTOM_CSS = """
<style>
/* Overall app background – purple/violet/blue gradient */
.stApp {
    background: linear-gradient(135deg, #001F3F 0%, #0074D9 50%, #7FDBFF 100%);
    color: #FFFFFF;
}

/* Make all main-area text white by default, but keep form inputs readable */
[data-testid="stAppViewContainer"] *:not(input):not(textarea):not(select) {
    color: #FFFFFF;
}

/* Make main content markdown text white & nicely styled */
.stMarkdown, .stMarkdown p {
    font-family: "Segoe UI", system-ui, -apple-system, BlinkMacSystemFont, sans-serif;
    font-size: 16px;
    line-height: 1.6;
}

/* Ensure sidebar text stays light background but dark body text */
[data-testid="stSidebar"] {
    background-color: rgba(255, 255, 255, 0.95);
    backdrop-filter: blur(12px);
    border-right: 1px solid rgba(148, 163, 184, 0.45);
    color: #0f172a;
}
[data-testid="stSidebar"] * {
    color: #0f172a !important;
}

/* Sidebar inner padding */
[data-testid="stSidebar"] > div {
    padding-top: 1rem;
}

/* Sidebar headers */
[data-testid="stSidebar"] h1,
[data-testid="stSidebar"] h2,
[data-testid="stSidebar"] h3 {
    color: #0f172a;
    font-weight: 700;
}

/* Navigation radio group styled as soft pills */
[data-testid="stSidebar"] .stRadio > label {
    font-weight: 600;
}

[data-testid="stSidebar"] .stRadio div[role="radiogroup"] > label {
    border-radius: 999px;
    padding: 0.25rem 0.75rem;
    margin-bottom: 0.15rem;
}

/* Make headings (main area) nicely weighted & white */
.stMarkdown h1, .stMarkdown h2, .stMarkdown h3,
.stMarkdown h4, .stMarkdown h5, .stMarkdown h6 {
    font-weight: 600;
    margin-top: 0.75rem;
    margin-bottom: 0.4rem;
}

/* Nicer spacing for bullets */
.stMarkdown ul, .stMarkdown ol {
    padding-left: 1.4rem;
    margin-bottom: 0.75rem;
}

/* Global gradient buttons with white border & text */
.stButton > button {
    background: linear-gradient(135deg, #001F3F 0%, #0074D9 50%, #7FDBFF 100%);
    color: #FFFFFF !important;
    border: 1px solid #FFFFFF !important;
    border-radius: 8px;
    padding: 0.5rem 1.25rem;
    font-weight: 600;
}

/* Slight hover effect for buttons */
.stButton > button:hover {
    filter: brightness(1.05);
    box-shadow: 0 0 0.5rem rgba(15, 23, 42, 0.35);
}

/* Sidebar-specific button override: keep text WHITE */
[data-testid="stSidebar"] .stButton > button {
    color: #FFFFFF !important;
    background: linear-gradient(135deg, #000080, #312E81);
    border: 1px solid #000080 !important;
}

/* Also force any inner spans/icons inside sidebar buttons to be white */
[data-testid="stSidebar"] .stButton > button * {
    color: #FFFFFF !important;
}
</style>
"""
st.markdown(CUSTOM_CSS, unsafe_allow_html=True)


# =========================================================
#   SIMPLE & ROBUST LOGIN (MULTI-USER, SESSION = USERNAME)
# =========================================================
def check_password():
    """Returns True if the user entered valid credentials (multi-user support)."""

    auth = st.secrets.get("auth", {})

    if not auth:
        st.error(
            "🔐 Authentication is not configured.\n\n"
            "Please add user credentials to this app's Secrets in Streamlit Cloud, e.g.:\n\n"
            "[auth]\n"
            'AhsanKamal = "1234567"\n'
            'SalmanAgha = "123232"\n'
            'KunzaRizvi = "123222"\n'
        )
        return False

    def password_entered():
        entered_username = st.session_state.get("username", "")
        entered_password = st.session_state.get("password", "")

        if entered_username in auth and auth[entered_username] == entered_password:
            st.session_state["password_correct"] = True
            st.session_state["current_user"] = entered_username

            if "user_sessions" not in st.session_state:
                st.session_state["user_sessions"] = {}

            if entered_username not in st.session_state["user_sessions"]:
                st.session_state["user_sessions"][entered_username] = {
                    "tender_filename": None,
                    "tender_clauses": [],
                    "tenders": [],
                    "history": [],
                }

            st.session_state.pop("password", None)
        else:
            st.session_state["password_correct"] = False

    if "password_correct" not in st.session_state:
        st.markdown("## 🐼 Welcome To Wavetec's Tender Response Tool !")
        st.text_input("Username", key="username")
        st.text_input("Password", type="password", key="password")
        if st.button("Login!"):
            password_entered()
        return False

    if not st.session_state["password_correct"]:
        st.markdown("## 🐼 Welcome To Wavetec's Tender Response Tool !")
        st.text_input("Username", key="username")
        st.text_input("Password", type="password", key="password")
        if st.button("Login!"):
            password_entered()
        st.error("❌ Incorrect username or password")
        return False

    return True


def get_current_user():
    return st.session_state.get("current_user")


def get_user_session():
    user = get_current_user()
    if not user:
        return None

    if "user_sessions" not in st.session_state:
        st.session_state["user_sessions"] = {}

    if user not in st.session_state["user_sessions"]:
        st.session_state["user_sessions"][user] = {
            "tender_filename": None,
            "tender_clauses": [],
            "tenders": [],
            "history": [],
        }

    if "tenders" not in st.session_state["user_sessions"][user]:
        st.session_state["user_sessions"][user]["tenders"] = []

    return st.session_state["user_sessions"][user]


def logout():
    user = st.session_state.get("current_user")

    if user and "user_sessions" in st.session_state:
        st.session_state["user_sessions"].pop(user, None)

    st.session_state.pop("current_user", None)
    st.session_state.pop("password_correct", None)

    st.rerun()


# ---------- STOP APP IF NOT LOGGED IN ----------
if not check_password():
    st.stop()

current_user = get_current_user()
user_session = get_user_session()

# ---------- TITLE & INTRO ----------
st.title("🐼 Wavetec Tender Response Automation Tool !")
st.write("Central knowledge base for automated tender and RFP responses.")
if current_user:
    st.caption(f"Logged in as **{current_user}**")


# =========================================================
#   PATHS
# =========================================================
BASE_DIR = Path(__file__).parent

LIBRARY_ROOT = BASE_DIR / "Tender_Aligned_FinalLibrary"

CATEGORY_FOLDERS = {
    "Corporate Profile": LIBRARY_ROOT / "Corporate Profile",
    "Technical Profile": LIBRARY_ROOT / "Technical Profile",
    "Security Profile": LIBRARY_ROOT / "Security Profile",
    "Services And Delivery": LIBRARY_ROOT / "Services And Delivery",
}


# =========================================================
#   HELPER: READ WORD DOCUMENT
# =========================================================
def load_docx_text(file_path: Path) -> str:
    if not file_path.exists():
        return "❗ This document does not exist: " + str(file_path)
    doc = Document(file_path)
    paragraphs = [p.text for p in doc.paragraphs]
    return "\n\n".join(p for p in paragraphs if p.strip())


def list_docx_files(folder: Path):
    if not folder.exists():
        return []
    return sorted([p for p in folder.glob("*.docx")])


# =========================================================
#   OPENAI CLIENT & HELPERS
# =========================================================
def get_openai_client():
    if not OPENAI_AVAILABLE or OpenAI is None:
        st.error("❌ OpenAI Python client is not installed. Check requirements.txt for 'openai'.")
        return None

    api_key = st.secrets.get("OPENAI_API_KEY")
    if not api_key and "openai" in st.secrets:
        api_key = st.secrets["openai"].get("api_key")

    if not api_key:
        st.error("❌ OPENAI_API_KEY not found in secrets. Please add it in Streamlit Cloud settings.")
        return None

    try:
        client = OpenAI(api_key=api_key)
        return client
    except Exception as e:
        st.error(f"❌ Failed to create OpenAI client: {e}")
        return None


def embed_texts_with_openai(texts, client=None):
    if client is None:
        client = get_openai_client()
    if client is None:
        return None

    try:
        resp = client.embeddings.create(
            model="text-embedding-3-small",
            input=[t[:7000] for t in texts],
        )
        embeddings = [item.embedding for item in resp.data]
        return embeddings
    except Exception as e:
        st.warning(f"⚠️ Could not generate embeddings: {e}")
        return None


def generate_model_response(prompt_text: str):
    client = get_openai_client()
    if client is None:
        return None

    try:
        resp = client.responses.create(
            model="gpt-4.1-mini",
            input=prompt_text,
            temperature=0.2,
        )

        pieces = []
        for part in resp.output[0].content:
            if getattr(part, "type", None) == "output_text":
                pieces.append(part.text)

        return "\n".join(pieces) if pieces else None

    except Exception as e:
        st.error(f"❌ OpenAI model call failed: {e}")
        return None


def strip_code_fences(text: str) -> str:
    if not text:
        return text
    text = text.strip()
    if text.startswith("```"):
        lines = text.splitlines()
        if lines and lines[0].lstrip().startswith("```"):
            lines = lines[1:]
        if lines and lines[-1].lstrip().startswith("```"):
            lines = lines[:-1]
        text = "\n".join(lines).strip()
    return text


# =========================================================
#   SIMPLE INDEXING (WITH EMBEDDINGS)
# =========================================================
def index_library():
    index = []
    texts_for_embedding = []
    entries_without_embeddings = []

    for category, folder in CATEGORY_FOLDERS.items():
        doc_files = list_docx_files(folder)
        for doc_path in doc_files:
            text = load_docx_text(doc_path)
            if not text.strip():
                continue

            entry = {
                "category": category,
                "file_name": doc_path.name,
                "file_path": str(doc_path),
                "text": text,
                "embedding": None,
            }

            index.append(entry)
            texts_for_embedding.append(text)
            entries_without_embeddings.append(entry)

    if texts_for_embedding:
        client = get_openai_client()
        embeddings = embed_texts_with_openai(texts_for_embedding, client=client)
        if embeddings is not None:
            for entry, emb in zip(entries_without_embeddings, embeddings):
                entry["embedding"] = emb
            st.success("✅ Embeddings generated for indexed documents.")
        else:
            st.warning("⚠️ Proceeding without embeddings (keyword search only).")

    st.session_state["library_index"] = index
    return index


# =========================================================
#   CLAUSE ROUTING
# =========================================================
def route_clause_to_filters(clause_no: str, clause_text: str):
    text = f"{clause_no} {clause_text}".lower()

    allowed_categories = None
    filename_keywords = []

    if any(
        phrase in text
        for phrase in [
            "resume", "résumé", "curriculum vitae", "curriculum-vitae",
            "key personnel", "key staff", "key resources", "cv ", " c.v."
        ]
    ):
        allowed_categories = ["Corporate Profile"]
        filename_keywords = [
            "resume", "resumes", "cv", "curriculum", "key_personnel",
            "key-personnel", "key_staff", "key-staff"
        ]

    if any(
        phrase in text
        for phrase in [
            "organisation structure", "organizational structure",
            "organizational chart", "organisation chart",
            "organisational structure", "organisational chart",
            "organogram", "organigram", "org chart"
        ]
    ):
        allowed_categories = ["Corporate Profile"]
        filename_keywords = [
            "org", "organogram", "organigram", "organization",
            "organisation", "org_chart", "org-chart", "structure"
        ]

    return allowed_categories, filename_keywords


# =========================================================
#   RETRIEVAL (EMBEDDINGS + ROUTING)
# =========================================================
def cosine_similarity(a, b):
    dot = sum(x * y for x, y in zip(a, b))
    norm_a = math.sqrt(sum(x * x for x in a))
    norm_b = math.sqrt(sum(y * y for y in b))
    if norm_a == 0 or norm_b == 0:
        return 0.0
    return dot / norm_a


def get_relevant_library_entries(clause_text: str, clause_no: str = "", top_k: int = 5):
    library_index = st.session_state.get("library_index", [])
    if not library_index:
        return []

    words = re.findall(r"\w+", clause_text.lower())
    keywords = {w for w in words if len(w) > 3}

    allowed_categories, filename_keywords = route_clause_to_filters(clause_no, clause_text)

    clause_embedding = None
    if any(entry.get("embedding") is not None for entry in library_index):
        emb_list = embed_texts_with_openai([clause_text])
        if emb_list is not None:
            clause_embedding = emb_list[0]

    scored = []
    for entry in library_index:
        if allowed_categories and entry["category"] not in allowed_categories:
            continue

        text_lower = entry["text"].lower()

        if clause_embedding is not None and entry.get("embedding") is not None:
            sim = cosine_similarity(clause_embedding, entry["embedding"])
            base_score = sim * 100
        else:
            base_score = sum(1 for w in keywords if w in text_lower)

        filename_score = 0
        if filename_keywords:
            name_lower = entry["file_name"].lower()
            filename_score = sum(1 for k in filename_keywords if k in name_lower) * 10

        score = base_score + filename_score

        if score > 0:
            scored.append((score, entry))

    if not scored:
        for entry in library_index:
            text_lower = entry["text"].lower()
            base_score = sum(1 for w in keywords if w in text_lower)
            if base_score > 0:
                scored.append((base_score, entry))

    scored.sort(key=lambda x: x[0], reverse=True)
    top_entries = [e for score, e in scored[:top_k]]
    return top_entries


# =========================================================
#   HELPERS FOR TENDER UPLOAD / CLAUSE EXTRACTION
# =========================================================
def extract_text_from_pdf(uploaded_file) -> str:
    if not PDF_SUPPORT or pdfplumber is None:
        st.error("PDF support is not available on this deployment. Please upload DOCX or Excel instead.")
        return ""
    all_text = []
    with pdfplumber.open(uploaded_file) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                all_text.append(page_text)
    return "\n".join(all_text)


def extract_text_from_docx_file(uploaded_file) -> str:
    doc = Document(uploaded_file)
    paragraphs = [p.text for p in doc.paragraphs]
    return "\n".join(p for p in paragraphs if p.strip())


def extract_text_from_excel_file(uploaded_file) -> str:
    try:
        sheets = pd.read_excel(uploaded_file, sheet_name=None, header=None)
    except Exception as e:
        st.error(f"❌ Could not read Excel file: {e}")
        return ""

    parts = []
    for sheet_name, df in sheets.items():
        parts.append(f"Sheet: {sheet_name}")
        for _, row in df.iterrows():
            for cell in row:
                if pd.notna(cell):
                    parts.append(str(cell))
    return "\n".join(parts)


def looks_like_clause_id(value) -> bool:
    if pd.isna(value):
        return False
    s = str(value).strip()
    if not s:
        return False

    if re.match(r'^\d+(?:\.\d+)*\.?$', s):
        return True
    if re.match(r'^[A-Za-z]\.?$', s):
        return True
    if re.match(r'^[ivxlcdmIVXLCDM]+\.?$', s):
        return True

    return False


def extract_clauses_from_excel(uploaded_file):
    try:
        sheets = pd.read_excel(uploaded_file, sheet_name=None, header=None)
    except Exception as e:
        st.error(f"❌ Could not read Excel file for clause extraction: {e}")
        return []

    clauses = []
    for sheet_name, df in sheets.items():
        for _, row in df.iterrows():
            if row.isna().all():
                continue

            first = row.iloc[0] if len(row) > 0 else None
            second = row.iloc[1] if len(row) > 1 else None

            if looks_like_clause_id(first) and pd.notna(second) and str(second).strip():
                clause_no = str(first).strip()
                clause_text = str(second).strip()
                clauses.append({"clause_no": clause_no, "clause_text": clause_text})
    return clauses


def extract_clauses_from_text(raw_text: str):
    lines = [line.strip() for line in raw_text.splitlines()]

    clause_pattern = re.compile(
        r'^('
        r'\d+(?:\.\d+)*\.?'
        r'|[A-Za-z]\.?'
        r'|[ivxlcdm]+\b\.?'
        r')\s+(.*)$',
        re.IGNORECASE
    )

    clauses = []
    current_clause = None

    for line in lines:
        if not line:
            continue

        match = clause_pattern.match(line)
        if match:
            if current_clause:
                clauses.append(current_clause)

            clause_no = match.group(1).strip()
            clause_text = match.group(2).strip()
            current_clause = {"clause_no": clause_no, "clause_text": clause_text}
        else:
            if current_clause:
                current_clause["clause_text"] += " " + line

    if current_clause:
        clauses.append(current_clause)

    return clauses


# =========================================================
#   SIDEBAR (USER + LOGOUT + NAVIGATION)
# =========================================================
with st.sidebar:
    if current_user:
        st.markdown(f"**User:** {current_user}")
        if st.button("🚪 Logout"):
            logout()

    st.header("🧭 Navigation")
    page = st.radio(
        "Go to:",
        [
            "📖 View Documents",
            "🧠 Prepare / Index Library",
            "🧬 Upload Tender & Extract Clauses",
            "🧙 Generate Responses",
            "🗂 History",
        ],
        index=0
    )


# =========================================================
#   PAGE: PREPARE / INDEX LIBRARY
# =========================================================
if page == "🧠 Prepare / Index Library":
    st.subheader("🧠 Prepare / Index Wavetec Library")

    st.markdown(
        """
        This page prepares your Wavetec Tender Library for automated RFP responses.

        What this button will do:
        - Scan all categories
        - Read their full text
        - Build an in-memory index
        - Generate embeddings where possible
        - Show you how many documents were indexed
        """
    )

    if st.button("🚀 Index / Refresh Library"):
        with st.spinner("Indexing library (and generating embeddings)..."):
            index = index_library()
        st.success(f"✅ Indexed {len(index)} document entries into memory.")

        if index:
            st.markdown("### 📄 Indexed Documents (Preview)")
            preview = [{"Category": item["category"], "File": item["file_name"]} for item in index]
            st.dataframe(preview)


# =========================================================
#   PAGE: UPLOAD TENDER & EXTRACT CLAUSES
# =========================================================
elif page == "🧬 Upload Tender & Extract Clauses":
    st.subheader("🧬 Upload Tender & Extract Clauses")

    st.markdown(
        """
        Upload tender file(s).

        The app will:
        - Extract the full text
        - Detect clause numbers
        - Build a Tender Response Map (clause number + clause text)
        - Store it for later use (per logged-in user).
        """
    )

    base_types = ["docx", "xlsx", "xls"]
    if PDF_SUPPORT:
        allowed_types = base_types + ["pdf"]
        st.info("📎 You can upload PDF, Word (.docx), or Excel (.xlsx/.xls) tenders.")
    else:
        allowed_types = base_types
        st.warning("⚠️ PDF support is not available on this deployment. Please upload Word (.docx) or Excel (.xlsx/.xls).")

    uploaded_files = st.file_uploader(
        "Upload Tender Document(s) (up to 2)",
        type=allowed_types,
        accept_multiple_files=True,
        help="Accepted formats: " + ", ".join(f".{ext}" for ext in allowed_types) + ". You can upload up to 2 tenders."
    )

    if uploaded_files:
        if len(uploaded_files) > 2:
            st.warning("⚠️ You uploaded more than 2 files. Only the first 2 will be processed.")
            uploaded_files = uploaded_files[:2]

        for f in uploaded_files:
            st.info(f"📁 File uploaded: **{f.name}**")

        if st.button("🔍 Extract Clauses"):
            with st.spinner("Extracting text and detecting clauses for all uploaded tenders..."):
                extracted_tenders = []
                combined_clauses = []
                combined_filenames = []

                for uploaded_file in uploaded_files:
                    filename = uploaded_file.name.lower()
                    raw_text = ""
                    clauses = []
                    is_excel = False

                    if filename.endswith(".pdf"):
                        raw_text = extract_text_from_pdf(uploaded_file)
                    elif filename.endswith(".docx"):
                        raw_text = extract_text_from_docx_file(uploaded_file)
                    elif filename.endswith(".xlsx") or filename.endswith(".xls"):
                        raw_text = extract_text_from_excel_file(uploaded_file)
                        is_excel = True
                    else:
                        st.error(f"❌ Unsupported file type: {uploaded_file.name}")
                        raw_text = ""

                    if is_excel:
                        clauses = extract_clauses_from_excel(uploaded_file)
                        if not clauses:
                            st.error(
                                f"❌ No clauses could be extracted from Excel file: {uploaded_file.name}. "
                                "Check that column 1 contains clause IDs and column 2 contains clause text."
                            )
                    else:
                        if not raw_text or not raw_text.strip():
                            st.error(f"❌ No text could be extracted from this file: {uploaded_file.name}.")
                        else:
                            clauses = extract_clauses_from_text(raw_text)
                            if not clauses:
                                st.warning(
                                    f"⚠️ No clauses were detected in: {uploaded_file.name}. "
                                    "Check if the document uses recognizable numbering like '3.14', '4.2.1', 'A.', 'i.' at the start of lines."
                                )

                    if clauses:
                        combined_filenames.append(uploaded_file.name)

                        tagged_clauses = []
                        for c in clauses:
                            cc = dict(c)
                            cc["tender_filename"] = uploaded_file.name
                            tagged_clauses.append(cc)

                        extracted_tenders.append({
                            "tender_filename": uploaded_file.name,
                            "tender_clauses": tagged_clauses,
                        })

                        combined_clauses.extend(tagged_clauses)

                if extracted_tenders:
                    user_session = get_user_session()
                    if user_session is not None:
                        user_session["tenders"] = extracted_tenders
                        user_session["tender_clauses"] = combined_clauses
                        user_session["tender_filename"] = ", ".join(combined_filenames) if combined_filenames else None

                        if "history" not in user_session:
                            user_session["history"] = []
                        user_session["history"].append({
                            "timestamp": datetime.utcnow().isoformat(),
                            "type": "upload",
                            "tender_filename": user_session.get("tender_filename"),
                            "clauses": [c["clause_no"] for c in combined_clauses],
                            "response_text": None,
                        })

                    st.success(f"✅ Extracted {len(combined_clauses)} clauses across {len(extracted_tenders)} tender(s).")

                    for t in extracted_tenders:
                        st.markdown(f"### 📋 Tender Response Map (Preview) — **{t['tender_filename']}**")
                        df_clauses = pd.DataFrame(
                            [{"clause_no": c["clause_no"], "clause_text": c["clause_text"]} for c in t["tender_clauses"]]
                        )
                        st.dataframe(df_clauses, use_container_width=True)

                        with st.expander(f"🔎 View first 5 clauses (full text) — {t['tender_filename']}"):
                            for row in t["tender_clauses"][:5]:
                                st.markdown(f"**Clause {row['clause_no']}**")
                                st.write(row["clause_text"])
                                st.markdown("---")
                else:
                    st.error("❌ No clauses were extracted from the uploaded file(s).")
    else:
        st.info("📥 Please upload tender file(s) to begin.")


# =========================================================
#   PAGE: 🧙 GENERATE RESPONSES
# =========================================================
elif page == "🧙 Generate Responses":
    st.subheader("🧙 Generate Clause-by-Clause Responses")

    st.markdown(
        """
        This page helps you generate Wavetec responses for each tender clause.

        Workflow:
        1. Index the Wavetec library on the "🧠 Prepare / Index Library" page.
        2. Extract clauses on the "🧬 Upload Tender & Extract Clauses" page.
        3. Select one or more clauses.
        4. The app finds relevant library entries using embeddings (falls back to keyword search).
        5. Generate either per-clause responses or one unified response for all selected clauses.
        """
    )

    user_session = get_user_session()
    if user_session is None:
        st.error("No active user session found.")
    else:
        tender_clauses = user_session.get("tender_clauses", [])
        library_index = st.session_state.get("library_index")

        if not tender_clauses:
            st.warning("⚠️ No tender clauses found for this user. Please upload a tender and run Extract Clauses.")
        elif not library_index:
            st.warning("⚠️ Library index is empty. Please index your documents first.")
        else:
            with st.expander("📝 Or type a clause text manually (optional)"):
                manual_clause_no = st.text_input("Manual clause number (optional)", value="")
                manual_clause_text = st.text_area(
                    "Manual clause text (optional)",
                    value="",
                    height=140,
                    help="If filled, the app will also generate based on this text (even if you select no extracted clauses)."
                )

            options = []
            for c in tender_clauses:
                snippet = c["clause_text"][:80].replace("\n", " ")
                tname = c.get("tender_filename", "Tender")
                options.append(f"[{tname}] {c['clause_no']} – {snippet}...")

            selected_labels = st.multiselect(
                "Select one or more clauses to respond to:",
                options,
            )

            selected_clauses = []

            if selected_labels:
                selected_indices = [options.index(label) for label in selected_labels]
                selected_clauses.extend([tender_clauses[i] for i in selected_indices])

            if manual_clause_text and manual_clause_text.strip():
                selected_clauses.append({
                    "clause_no": manual_clause_no.strip() if manual_clause_no.strip() else "MANUAL",
                    "clause_text": manual_clause_text.strip(),
                    "tender_filename": "Manual Input",
                })

            if not selected_clauses:
                st.info("Select at least one clause above, and/or type a manual clause text to proceed.")
            else:
                st.markdown("---")
                for clause in selected_clauses:
                    tname = clause.get("tender_filename", None)
                    if tname:
                        st.markdown(f"### 📌 Clause {clause['clause_no']}  _(from: {tname})_")
                    else:
                        st.markdown(f"### 📌 Clause {clause['clause_no']}")
                    st.write(clause["clause_text"])
                    st.markdown("")

                all_relevant_entries = []
                with st.spinner("Finding relevant Wavetec library content for all selected clauses..."):
                    for clause in selected_clauses:
                        entries = get_relevant_library_entries(
                            clause["clause_text"],
                            clause_no=clause.get("clause_no", ""),
                            top_k=5,
                        )
                        all_relevant_entries.append(entries)

                for clause, entries in zip(selected_clauses, all_relevant_entries):
                    st.markdown(f"### 📚 Top Matching Library Entries for Clause {clause['clause_no']}")
                    if not entries:
                        st.write("_No relevant library entries were found for this clause._")
                    else:
                        for i, entry in enumerate(entries, start=1):
                            st.markdown(f"**{i}. {entry['category']} – {entry['file_name']}**")
                            snippet = entry["text"][:400].replace("\n", " ")
                            st.write(snippet + "...")
                            st.markdown("---")

                per_clause_context_blocks = []
                for entries in all_relevant_entries:
                    context_blocks = []
                    for entry in entries:
                        block = (
                            f"Category: {entry['category']}\n"
                            f"File: {entry['file_name']}\n"
                            f"Content:\n{entry['text']}\n"
                            "-------------------------\n"
                        )
                        context_blocks.append(block)
                    per_clause_context_blocks.append("\n".join(context_blocks))

                prompt_header = """You are the Wavetec RFP Response Engine.

Use ONLY the library context below to answer the clauses. Do NOT invent facts that are not supported by the context.

"""

                clause_sections = []
                for idx, (clause, ctx) in enumerate(zip(selected_clauses, per_clause_context_blocks), start=1):
                    section = f"""=== Clause {idx} ===
Clause number: {clause['clause_no']}
Clause text:
\"\"\"{clause['clause_text']}\"\"\"


Wavetec Library Context:
\"\"\"{ctx}\"\"\"


"""
                    clause_sections.append(section)

                prompt_footer = """TASK:
For each clause above:

1. Write a detailed, structured, bid-winning response from Wavetec's perspective.
2. Use a formal, government/RFP-appropriate tone.
3. Include all relevant technical, architectural, security, delivery, and operational details you can find in the context.
4. After the main response for each clause, add:

**Images/Description**
- describe the recommended image(s) in detail (what they should show, why they are useful, and where they would appear in the final document). If no image is needed, you may skip this.

**Gaps / Missing Information:**
- Bullet list of items where the library does not contain enough detail.

**Assumptions:**
- Bullet list of assumptions you are making to answer this clause.

Return the answer in markdown format, and do NOT wrap it inside ``` code fences.
"""

                prompt_text_per_clause = prompt_header + "\n".join(clause_sections) + prompt_footer

                combined_numbers = ", ".join(c["clause_no"] for c in selected_clauses)
                combined_clause_lines = [f"{c['clause_no']}: {c['clause_text']}" for c in selected_clauses]
                combined_clause_text = "\n".join(combined_clause_lines)
                combined_context = "\n".join(per_clause_context_blocks)

                concat_prompt_text = f"""You are the Wavetec RFP Response Engine.

Use ONLY the library context below to answer the combined requirement. Do NOT invent facts that are not supported by the context.

Combined clauses: {combined_numbers}

Combined clause text (treat these as one integrated requirement):
\"\"\"{combined_clause_text}\"\"\"


Wavetec Library Context (for all subclauses combined):
\"\"\"{combined_context}\"\"\"


TASK:
1. Write a detailed, structured, bid-winning response from Wavetec's perspective.
2. Use a formal, government/RFP-appropriate tone.
3. Include all relevant technical, architectural, security, delivery, and operational details you can find in the context.
4. Decide the appropriate length and depth of each response yourself, based on how complex and critical the clause is.
5. After the main response for each clause, add:

**Images/Description**
- describe the recommended image(s) in detail (what they should show, why they are useful, and where they would appear in the final document). If no image is needed, you may skip this.

**Gaps / Missing Information:**
- Bullet list of items where the library does not contain enough detail.

**Assumptions:**
- Bullet list of assumptions you are making to answer this clause.

Return the answer in markdown format, and do NOT wrap it inside ``` code fences.
"""

                st.markdown("### ✍️ Prepared Prompt – per clause")
                st.text_area(
                    "Prompt for ChatGPT / LLM (per-clause responses)",
                    value=prompt_text_per_clause,
                    height=260,
                )

                st.markdown("### ✍️ Prepared Prompt – concatenated (single unified response)")
                st.text_area(
                    "Prompt for ChatGPT / LLM (concatenated clauses)",
                    value=concat_prompt_text,
                    height=260,
                )

                col1, col2 = st.columns(2)

                with col1:
                    if st.button("🧙 Generate Responses (per clause)"):
                        with st.spinner("Calling OpenAI to generate responses for all selected clauses..."):
                            answer = generate_model_response(prompt_text_per_clause)

                        if answer:
                            answer_clean = strip_code_fences(answer)
                            st.markdown("### ✅ Generated Responses (per clause)")
                            with st.container(border=True):
                                st.markdown(answer_clean)

                            if user_session is not None:
                                if "history" not in user_session:
                                    user_session["history"] = []
                                user_session["history"].append({
                                    "timestamp": datetime.utcnow().isoformat(),
                                    "type": "per_clause",
                                    "tender_filename": user_session.get("tender_filename"),
                                    "clauses": [c["clause_no"] for c in selected_clauses],
                                    "response_text": answer_clean,
                                })
                        else:
                            st.warning("No response was returned from the model (per-clause mode).")

                with col2:
                    if st.button("🧙‍♂️ Generate Single Response (concatenated)"):
                        with st.spinner("Calling OpenAI to generate a unified response for all selected clauses..."):
                            concat_answer = generate_model_response(concat_prompt_text)

                        if concat_answer:
                            concat_answer_clean = strip_code_fences(concat_answer)
                            st.markdown("### ✅ Generated Response (concatenated clauses)")
                            with st.container(border=True):
                                st.markdown(concat_answer_clean)

                            if user_session is not None:
                                if "history" not in user_session:
                                    user_session["history"] = []
                                user_session["history"].append({
                                    "timestamp": datetime.utcnow().isoformat(),
                                    "type": "concatenated",
                                    "tender_filename": user_session.get("tender_filename"),
                                    "clauses": [c["clause_no"] for c in selected_clauses],
                                    "response_text": concat_answer_clean,
                                })
                        else:
                            st.warning("No response was returned from the model (concatenated mode).")


# =========================================================
#   PAGE: VIEW DOCUMENTS
# =========================================================
elif page == "📖 View Documents":
    st.sidebar.header("📂 Document Library")

    category = st.sidebar.selectbox(
        "Select a category:",
        list(CATEGORY_FOLDERS.keys())
    )

    folder_path = CATEGORY_FOLDERS[category]
    doc_files = list_docx_files(folder_path)

    if not doc_files:
        st.warning(f"No .docx files found in folder: {folder_path.name}")
    else:
        doc_display_names = [f.name for f in doc_files]
        selected_doc_name = st.sidebar.selectbox("Select a document:", doc_display_names)

        selected_doc_path = folder_path / selected_doc_name

        st.subheader(f"{category} → {selected_doc_name}")
        st.caption(f"Source file: {selected_doc_path.relative_to(BASE_DIR)}")

        content = load_docx_text(selected_doc_path)
        st.markdown(content)


# =========================================================
#   PAGE: 🗂 HISTORY (PER-USER)
# =========================================================
elif page == "🗂 History":
    st.subheader("🗂 Session History")

    user_session = get_user_session()
    if user_session is None:
        st.error("No active user session found.")
    else:
        history = user_session.get("history", [])

        if not history:
            st.info("No history found for this user yet. Generate some responses first.")
        else:
            if st.button("🧹 Clear all history for this user"):
                user_session["history"] = []
                st.success("All history for this user has been cleared.")
                st.rerun()

            st.markdown("### 🔎 History Entries (newest first)")

            for idx, entry in reversed(list(enumerate(history))):
                entry_type = entry.get("type", "unknown")
                tender_name = entry.get("tender_filename") or "N/A"
                timestamp = entry.get("timestamp", "N/A")
                clauses = entry.get("clauses", [])
                response_text = entry.get("response_text")

                if entry_type == "upload":
                    title = f"📁 Uploaded tender **{tender_name}** ({len(clauses)} clauses)"
                elif entry_type == "per_clause":
                    title = f"🧙 Per-clause response for tender **{tender_name}** (Clauses: {', '.join(clauses)})"
                elif entry_type == "concatenated":
                    title = f"🧙‍♂️ Concatenated response for tender **{tender_name}** (Clauses: {', '.join(clauses)})"
                else:
                    title = f"Entry type: {entry_type}"

                with st.expander(f"{title}  —  {timestamp}"):
                    st.write(f"**Tender:** {tender_name}")
                    st.write(f"**Type:** {entry_type}")
                    st.write(f"**Clauses:** {', '.join(clauses) if clauses else 'N/A'}")
                    st.write(f"**Timestamp (UTC):** {timestamp}")

                    if response_text:
                        st.markdown("---")
                        st.markdown("#### 📝 Response Text")
                        st.markdown(response_text)

                    if st.button("Delete this entry", key=f"delete_entry_{idx}"):
                        del history[idx]
                        user_session["history"] = history
                        st.success("History entry deleted.")
                        st.rerun()
